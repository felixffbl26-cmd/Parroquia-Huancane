"""
Generador de Informe Impreso - Parroquia Santiago Apóstol de Huancané
Versión Mejorada: Gráficos Premium y Narrativa Automática
"""

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import numpy as np
import warnings
warnings.filterwarnings('ignore')

# --- CONFIGURACIÓN ESTÉTICA GLOBAL ---
plt.style.use('seaborn-v0_8-whitegrid')
# Paleta personalizada: Azul Parroquial, Dorado Litúrgico, Rojo Mártir, Verde Esperanza
COLORES = ['#2C3E50', '#F39C12', '#E74C3C', '#27AE60', '#8E44AD']
sns.set_palette(COLORES)

class GeneradorInforme:
    def __init__(self, datos_csv):
        self.df = pd.read_csv(datos_csv)
        self.df['fecha'] = pd.to_datetime(self.df['fecha'], errors='coerce')
        self.ruta_graficos = r"C:\Users\FELIX\.gemini\antigravity\scratch\parroquia_huancane\informes\graficos"
        
        import os
        os.makedirs(self.ruta_graficos, exist_ok=True)
        
    def _agregar_etiquetas(self, ax, formato='S/ {:,.0f}'):
        """Helper para agregar etiquetas de valor en las barras"""
        for p in ax.patches:
            val = p.get_height()
            if val > 0:
                ax.annotate(formato.format(val), 
                            (p.get_x() + p.get_width() / 2., val), 
                            ha = 'center', va = 'center', 
                            xytext = (0, 9), 
                            textcoords = 'offset points',
                            fontsize=10, fontweight='bold', color='#34495e')

    def generar_graficos(self):
        """Genera gráficos de alta calidad con explicaciones para ambos años"""
        print("Generando gráficos premium...")
        
        # Obtener años disponibles
        años_disponibles = sorted(self.df['año'].unique())
        print(f"  Años detectados: {años_disponibles}")
        
        # 1. Comparativa de Ingresos por Año
        plt.figure(figsize=(10, 6))
        ax = plt.gca()
        ingresos_año = self.df.groupby('año')['ingreso'].sum()
        
        # Usar colores diferentes para cada año
        colores_años = [COLORES[i % len(COLORES)] for i in range(len(ingresos_año))]
        bars = plt.bar(ingresos_año.index.astype(str), ingresos_año.values, color=colores_años, width=0.6, edgecolor='white', linewidth=2)
        
        plt.title('Comparativa de Ingresos Totales por Año', fontsize=16, fontweight='bold', pad=20, color='#2c3e50')
        plt.ylabel('Monto (S/)', fontsize=12)
        plt.xlabel('Año', fontsize=12)
        plt.grid(axis='y', alpha=0.3)
        self._agregar_etiquetas(ax)
        sns.despine(left=True)
        plt.tight_layout()
        plt.savefig(f"{self.ruta_graficos}/1_ingresos_comparativa.png", dpi=300, bbox_inches='tight')
        plt.close()
        
        # 2. Evolución Mensual Comparativa (ambos años si existen)
        plt.figure(figsize=(14, 7))
        meses_lbl = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep', 'Oct', 'Nov', 'Dic']
        
        for idx, año in enumerate(años_disponibles):
            df_año = self.df[self.df['año'] == año]
            if not df_año.empty:
                ingresos_mes = df_año.groupby('mes')['ingreso'].sum()
                # Rellenar meses faltantes con 0 para continuidad
                ingresos_mes = ingresos_mes.reindex(range(1, 13), fill_value=0)
                
                color = COLORES[idx % len(COLORES)]
                plt.plot(range(1, 13), ingresos_mes.values, marker='o', linewidth=3, 
                        color=color, markersize=8, label=f'Año {año}', alpha=0.8)
                plt.fill_between(range(1, 13), ingresos_mes.values, color=color, alpha=0.1)
        
        plt.title('Tendencia de Ingresos Mensuales - Comparativa', fontsize=16, fontweight='bold', pad=20, color='#2c3e50')
        plt.xticks(range(1, 13), meses_lbl)
        plt.ylabel('Ingresos (S/)', fontsize=12)
        plt.xlabel('Mes', fontsize=12)
        plt.legend(loc='best', fontsize=11, frameon=True, shadow=True)
        plt.grid(axis='both', alpha=0.3, linestyle='--')
        sns.despine()
        plt.tight_layout()
        plt.savefig(f"{self.ruta_graficos}/2_ingresos_mensuales_comparativa.png", dpi=300, bbox_inches='tight')
        plt.close()
        
        # 3. Ingresos por Categoría (Donut Chart Mejorado)
        plt.figure(figsize=(10, 8))
        ingresos_cat = self.df.groupby('categoria')['ingreso'].sum().sort_values(ascending=False)
        colors = sns.color_palette("pastel")
        
        # Donut Chart
        plt.pie(ingresos_cat.values, labels=ingresos_cat.index, autopct='%1.1f%%', 
                startangle=90, colors=colors, pctdistance=0.85,
                textprops={'fontsize': 11, 'fontweight': 'bold'})
        
        # Círculo blanco en el centro
        centre_circle = plt.Circle((0,0),0.70,fc='white')
        fig = plt.gcf()
        fig.gca().add_artist(centre_circle)
        
        plt.title('Distribución de Ingresos por Fuente', fontsize=16, fontweight='bold', pad=20, color='#2c3e50')
        plt.tight_layout()
        plt.savefig(f"{self.ruta_graficos}/3_ingresos_categoria.png", dpi=300)
        plt.close()
        
        # 4. Distribución de Sacramentos (Barra Horizontal)
        sacramentos = self.df[self.df['categoria'] == 'SACRAMENTO']
        if not sacramentos.empty:
            plt.figure(figsize=(10, 6))
            ax = plt.gca()
            cant_sacr = sacramentos.groupby('subcategoria').size().sort_values(ascending=True)
            
            colors_sacr = sns.color_palette("viridis", n_colors=len(cant_sacr))
            bars = plt.barh(cant_sacr.index, cant_sacr.values, color=colors_sacr)
            
            plt.title('Sacramentos Administrados', fontsize=16, fontweight='bold', pad=20, color='#2c3e50')
            plt.xlabel('Cantidad', fontsize=12)
            
            # Etiquetas en las barras
            for i, v in enumerate(cant_sacr.values):
                plt.text(v + 1, i, str(v), va='center', fontweight='bold')
                
            sns.despine(bottom=True, left=True)
            plt.grid(axis='x', alpha=0.3)
            plt.tight_layout()
            plt.savefig(f"{self.ruta_graficos}/4_sacramentos_distribucion.png", dpi=300)
            plt.close()

    def _generar_explicacion(self, metrica, df_filtrado, tipo='ingreso'):
        """Genera texto narrativo PROFUNDO e inteligente basado en los datos"""
        
        if tipo == 'comparativa_anual':
            ing_24 = df_filtrado[df_filtrado['año'] == 2024]['ingreso'].sum()
            ing_25 = df_filtrado[df_filtrado['año'] == 2025]['ingreso'].sum()
            diff = ing_25 - ing_24
            pct = (diff / ing_24 * 100) if ing_24 > 0 else 0
            
            txt = f"Al realizar el análisis comparativo financiero, se evidencia un {'crecimiento' if diff >= 0 else 'decrecimiento'} de las recaudaciones. "
            txt += f"El año 2025 cerró con S/ {ing_25:,.2f}, frente a los S/ {ing_24:,.2f} del periodo anterior. "
            if diff > 0:
                txt += f"Esto representa una variación positiva del {pct:.1f}%. Este superávit sugiere una mayor participación de los fieles o una gestión más eficiente de los eventos parroquiales."
            else:
                txt += f"Esto representa una caída del {abs(pct):.1f}%. Se recomienda analizar si hubo menos celebraciones sacramentales o factores externos que afectaron la economía local."
            return txt
            
        elif tipo == 'tendencia_mensual':
            mensual = df_filtrado.groupby('mes')['ingreso'].sum()
            mes_max = mensual.idxmax()
            mes_min = mensual.idxmin()
            promedio = mensual.mean()
            
            meses = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
            
            txt = f"El comportamiento mensual muestra una estacionalidad marcada. El mes más fuerte fue **{meses[mes_max-1]}** con S/ {mensual.max():,.2f}, superando ampliamente el promedio mensual de S/ {promedio:,.2f}. "
            txt += f"Por otro lado, **{meses[mes_min-1]}** registró la menor actividad con S/ {mensual.min():,.2f}. "
            txt += "Esta fluctuación es típica en la administración parroquial y suele estar correlacionada con el calendario litúrgico (Semana Santa, Fiestas Patronales, Navidad)."
            return txt

        return "El gráfico muestra la distribución de los datos financieros procesados."

    def generar_word(self):
        """Genera el informe Word con diseño mejorado"""
        print("\nGenerando documento Word...")
        doc = Document()
        
        # --- ESTILOS ---
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # --- PORTADA ---
        section = doc.sections[0]
        section.top_margin = Inches(2)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('INFORME DE GESTIÓN PARROQUIAL')
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(44, 62, 80)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Parroquia Santiago Apóstol de Huancané')
        run.italic = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph('\n\n\n\n')
        
        # Tabla resumen en portada
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells = table.rows[0].cells
        
        total_ing = self.df['ingreso'].sum()
        total_sacr = len(self.df[self.df['categoria'] == 'SACRAMENTO'])
        
        def _set_cell(cell, label, val):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}\n")
            run.font.bold = True
            run.font.color.rgb = RGBColor(127, 140, 141)
            run = p.add_run(val)
            run.font.size = Pt(14)
            run.font.bold = True
            
        _set_cell(cells[0], "Ingresos Totales", f"S/ {total_ing:,.0f}")
        _set_cell(cells[1], "Sacramentos", f"{total_sacr}")
        _set_cell(cells[2], "Documentos", f"{len(self.df[self.df['categoria'] == 'DOCUMENTO'])}")
        
        doc.add_page_break()
        
        # --- CONTENIDO ---
        def add_header(text):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(41, 128, 185)
            p.space_after = Pt(12)
            
        def add_chart_section(title, image_path, explanation):
            add_header(title)
            doc.add_picture(image_path, width=Inches(6))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run("\nAnálisis: ")
            run.bold = True
            p.add_run(explanation)
            doc.add_paragraph("\n") # Espacio

        # 1. Finanzas
        add_chart_section("1. Comparativa Anual de Ingresos", 
                         f"{self.ruta_graficos}/1_ingresos_comparativa.png",
                         self._generar_explicacion(None, self.df, 'comparativa_anual'))
        
        doc.add_page_break()

        add_chart_section("2. Comportamiento Mensual Comparativo", 
                         f"{self.ruta_graficos}/2_ingresos_mensuales_comparativa.png",
                         "Este gráfico muestra la evolución mensual de los ingresos comparando todos los años disponibles. Las líneas permiten identificar patrones estacionales y tendencias de crecimiento o decrecimiento entre periodos.")

        add_chart_section("3. Fuentes de Ingreso", 
                         f"{self.ruta_graficos}/3_ingresos_categoria.png",
                         "El gráfico muestra la distribución porcentual de los ingresos. Las misas y los sacramentos constituyen la base principal de la economía parroquial, lo que resalta la importancia de la actividad litúrgica constante.")

        doc.add_page_break()

        add_chart_section("4. Vida Sacramental", 
                         f"{self.ruta_graficos}/4_sacramentos_distribucion.png",
                         "Detalle de los sacramentos administrados a los fieles. Este indicador refleja no solo la actividad administrativa, sino el crecimiento espiritual de la comunidad parroquial.")
        
        # --- CONCLUSIÓN FINAL ---
        add_header("5. Conclusiones Generales")
        p = doc.add_paragraph()
        run = p.add_run(f"""
Tras el análisis de los datos del período 2024-2025, se concluye que la Parroquia Santiago Apóstol mantiene una dinámica activa. 

La gestión financiera muestra un balance de S/ {(self.df['ingreso'].sum() - self.df['egreso'].sum()):,.2f}, lo cual permite cubrir las necesidades operativas.

Se recomienda continuar fomentando la participación de los fieles en las celebraciones comunitarias, ya que representan el pilar tanto espiritual como económico de la institución.
""")
        
        output_path = r"C:\Users\FELIX\.gemini\antigravity\scratch\parroquia_huancane\informes\Informe_Profesional_2025.docx"
        doc.save(output_path)
        print(f"✓ Informe generado: {output_path}")

# Ejecutar
if __name__ == "__main__":
    app = GeneradorInforme(r"C:\Users\FELIX\.gemini\antigravity\scratch\parroquia_huancane\datos_consolidados.csv")
    app.generar_graficos()
    app.generar_word()
